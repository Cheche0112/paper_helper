import asyncio
import os
import sqlite3
from datetime import datetime, timedelta
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import matplotlib.pyplot as plt
from collections import Counter, defaultdict
import re
import pandas as pd
from pathlib import Path
import random
import argparse
import numpy as np
import requests
from urllib.parse import urlparse
from tqdm import tqdm

from arxiv_crawler import ArxivScraper
from arxiv_time import next_arxiv_update_day
from paper import PaperDatabase, PaperExporter
from categories import parse_categories
from call_ds import update_analyse_database, chat_v3, load_pdf_content

class ArxivWorkflow:
    def __init__(self, 
                 output_dir="./output_workflow",
                 category_whitelist=["cs.CV", "cs.AI", "cs.LG", "cs.CL", "cs.IR", "cs.MA"],
                 category_blacklist=[],
                 optional_keywords=["LLM", "GPT", "AI", "language model", "deep learning", "neural network"],
                 trans_to=None,
                 proxy=None,
                 date_from=None,
                 date_until=None,
                 days_back=7,
                 analyze_pdf=False):
        """
        初始化工作流
        
        Args:
            output_dir: 输出目录
            category_whitelist: 分类白名单
            category_blacklist: 分类黑名单
            optional_keywords: 关键词
            trans_to: 翻译目标语言，None表示不翻译
            proxy: 代理服务器
            date_from: 开始日期，格式为"YYYY-MM-DD"，如果为None则使用当前日期减去days_back天
            date_until: 结束日期，格式为"YYYY-MM-DD"，如果为None则使用当前日期
            days_back: 如果date_from为None，则使用当前日期减去days_back天作为开始日期
            analyze_pdf: 是否使用大模型分析PDF内容
        """
        self.output_dir = output_dir
        self.category_whitelist = category_whitelist
        self.category_blacklist = category_blacklist
        self.optional_keywords = optional_keywords
        self.trans_to = trans_to
        self.proxy = proxy
        self.analyze_pdf = analyze_pdf
        
        # 创建输出目录
        os.makedirs(output_dir, exist_ok=True)
        
        # 初始化数据库
        self.db = PaperDatabase("papers.db")
        
        # 设置日期范围
        self.today = datetime.now()
        
        # 如果提供了具体日期，则使用提供的日期
        if date_until:
            self.date_until = date_until
        else:
            self.date_until = self.today.strftime("%Y-%m-%d")
            
        if date_from:
            self.date_from = date_from
        else:
            self.date_from = (self.today - timedelta(days=days_back)).strftime("%Y-%m-%d")
        
        # 初始化爬虫
        self.scraper = None
    
    def set_date_range(self, date_from, date_until):
        """
        设置自定义日期范围
        
        Args:
            date_from: 开始日期，格式为"YYYY-MM-DD"
            date_until: 结束日期，格式为"YYYY-MM-DD"
        """
        self.date_from = date_from
        self.date_until = date_until
        print(f"已设置日期范围: {self.date_from} 到 {self.date_until}")
        
    async def run(self):
        """运行完整工作流"""
        print(f"开始运行arXiv工作流...")
        print(f"日期范围: {self.date_from} 到 {self.date_until}")
        
        # 1. 爬取论文
        await self.crawl_papers()
        
        # 2. 分析论文
        analysis_results = self.analyze_papers()
        
        # 3. 阅读论文推荐
        reading_recommendations = self.reading_papers(analysis_results)
        
        # 4. 如果启用了PDF分析，则下载PDF并调用大模型分析PDF
        if self.analyze_pdf:
            print("准备下载和分析PDF内容...")
            self.download_pdfs(analysis_results["all_papers"])
            self.call_ds_analysis()
        
        # 5. 生成报告
        self.generate_report(analysis_results, reading_recommendations)
        
        print(f"工作流完成！结果已保存到 {self.output_dir} 目录")
    
    async def crawl_papers(self):
        """爬取论文"""
        print("开始爬取论文...")
        print(f"传递给爬虫的日期范围: {self.date_from} 到 {self.date_until}")
        
        # 初始化爬虫
        self.scraper = ArxivScraper(
            date_from=self.date_from,
            date_until=self.date_until,
            category_blacklist=self.category_blacklist,
            category_whitelist=self.category_whitelist,
            optional_keywords=self.optional_keywords,
            trans_to=self.trans_to,
            proxy=self.proxy
        )
        
        # 运行爬虫
        await self.scraper.fetch_all()
        print(f"爬取完成！共获取 {len(self.scraper.papers)} 篇论文")
        
        # 处理论文
        print("处理论文...")
        self.scraper.process_papers()
        
        # 导出为Markdown
        print("导出为Markdown...")
        self.scraper.paper_exporter.to_markdown(
            output_dir=os.path.join(self.output_dir, "markdown"), 
            filename_format="%Y-%m-%d",
            metadata=self.scraper.meta_data
        )
    
    def analyze_papers(self):
        """分析论文数据"""
        print("开始分析论文...")
        
        # 获取日期范围内的所有论文
        date_from = datetime.strptime(self.date_from, "%Y-%m-%d")
        date_until = datetime.strptime(self.date_until, "%Y-%m-%d")
        
        papers = []
        current_date = date_from
        while current_date <= date_until:
            papers.extend(self.db.fetch_papers_on_date(current_date))
            current_date += timedelta(days=1)
        
        print(f"分析范围内共有 {len(papers)} 篇论文")
        
        # 如果没有论文，返回空结果
        if not papers:
            return {
                "total_papers": 0,
                "date_distribution": {},
                "category_distribution": {},
                "top_keywords": [],
                "top_authors": [],
                "papers_by_category": {},
                "papers_by_date": {},
                "citation_network": {},
                "trending_topics": []
            }
        
        # 1. 按日期分布
        date_distribution = Counter([paper.first_announced_date.strftime("%Y-%m-%d") for paper in papers])
        
        # 2. 按分类分布
        category_counter = Counter()
        for paper in papers:
            for category in paper.categories:
                category_counter[category] += 1
        
        # 3. 提取关键词
        all_text = " ".join([f"{paper.title} {paper.abstract}" for paper in papers]).lower()
        # 简单的关键词提取，实际应用中可能需要更复杂的NLP方法
        words = re.findall(r'\b[a-z]{3,}\b', all_text)
        # 排除常见的停用词
        stopwords = {'the', 'and', 'of', 'to', 'a', 'in', 'for', 'is', 'on', 'that', 'by', 'this', 'with', 'are', 'be', 'as', 'an', 'we', 'our', 'from', 'such', 'or', 'can', 'has', 'have', 'not', 'which', 'these', 'their', 'was', 'were', 'been', 'also', 'but', 'than', 'more', 'its', 'they', 'when', 'who', 'will', 'would', 'could', 'should', 'may', 'might', 'must', 'one', 'two', 'three', 'first', 'second', 'third'}
        filtered_words = [word for word in words if word not in stopwords]
        top_keywords = Counter(filtered_words).most_common(20)
        
        # 4. 提取作者
        author_counter = Counter()
        for paper in papers:
            authors = paper.authors.split(',')
            for author in authors:
                author = author.strip()
                if author:
                    author_counter[author] += 1
        top_authors = author_counter.most_common(10)
        
        # 5. 按分类整理论文
        papers_by_category = defaultdict(list)
        for paper in papers:
            main_category = paper.categories[0]  # 使用第一个分类作为主分类
            papers_by_category[main_category].append(paper)
        
        # 6. 按日期整理论文
        papers_by_date = defaultdict(list)
        for paper in papers:
            date_str = paper.first_announced_date.strftime("%Y-%m-%d")
            papers_by_date[date_str].append(paper)
        
        # 7. 分析论文标题中的关键词趋势
        title_keywords = []
        for paper in papers:
            # 提取标题中的关键词
            title_words = re.findall(r'\b[a-zA-Z]{3,}\b', paper.title.lower())
            title_words = [word for word in title_words if word not in stopwords]
            title_keywords.extend(title_words)
        
        trending_topics = Counter(title_keywords).most_common(15)
        
        # 8. 尝试构建简单的引用网络（基于作者合作关系）
        citation_network = defaultdict(set)
        for paper in papers:
            authors = [author.strip() for author in paper.authors.split(',') if author.strip()]
            for author in authors:
                for co_author in authors:
                    if author != co_author:
                        citation_network[author].add(co_author)
        
        # 转换为可序列化的格式
        citation_network_serializable = {author: list(co_authors) for author, co_authors in citation_network.items()}
        
        # 返回分析结果
        return {
            "total_papers": len(papers),
            "date_distribution": dict(date_distribution),
            "category_distribution": dict(category_counter),
            "top_keywords": top_keywords,
            "top_authors": top_authors,
            "papers_by_category": papers_by_category,
            "papers_by_date": papers_by_date,
            "citation_network": citation_network_serializable,
            "trending_topics": trending_topics,
            "all_papers": papers  # 添加所有论文列表，供reading_papers使用
        }
    
    def reading_papers(self, analysis_results):
        """
        论文阅读推荐功能
        根据论文的引用次数、关键词匹配度和新颖性生成阅读推荐
        """
        print("生成论文阅读推荐...")
        
        if analysis_results["total_papers"] == 0:
            return {
                "must_read": [],
                "recommended": [],
                "interesting": [],
                "reading_path": []
            }
        
        all_papers = analysis_results.get("all_papers", [])
        
        # 1. 计算每篇论文的推荐分数
        # 这里使用一个简单的评分机制，实际应用中可以使用更复杂的算法
        paper_scores = []
        
        for paper in all_papers:
            # 基础分数
            score = 0
            
            # 根据标题中包含的热门关键词加分
            title_lower = paper.title.lower()
            for keyword, count in analysis_results["top_keywords"]:
                if keyword in title_lower:
                    score += min(count / 10, 2)  # 限制单个关键词的最大加分
            
            # 根据作者知名度加分
            authors = paper.authors.split(',')
            for author in authors:
                author = author.strip()
                for top_author, count in analysis_results["top_authors"]:
                    if author == top_author:
                        score += min(count / 2, 3)  # 限制单个作者的最大加分
            
            # 根据分类热度加分
            for category in paper.categories:
                if category in analysis_results["category_distribution"]:
                    cat_count = analysis_results["category_distribution"][category]
                    score += min(cat_count / 20, 1)  # 限制单个分类的最大加分
            
            # 添加一些随机性，模拟论文质量和创新性的评估
            # 实际应用中，这部分可以用更复杂的NLP模型或引用数据替代
            score += random.uniform(0, 3)
            
            paper_scores.append((paper, score))
        
        # 按分数排序
        paper_scores.sort(key=lambda x: x[1], reverse=True)
        
        # 2. 生成不同级别的阅读推荐
        must_read = [p for p, s in paper_scores[:min(5, len(paper_scores))]]
        recommended = [p for p, s in paper_scores[5:min(15, len(paper_scores))]]
        interesting = [p for p, s in paper_scores[15:min(30, len(paper_scores))]]
        
        # 3. 生成阅读路径
        # 这里简单地按主题分组，实际应用中可以构建更复杂的知识图谱
        reading_path = []
        
        # 尝试构建一个简单的阅读路径，从基础到高级
        if must_read:
            # 按分类分组
            path_by_category = defaultdict(list)
            for paper in must_read + recommended[:5]:
                main_category = paper.categories[0]
                path_by_category[main_category].append(paper)
            
            # 为每个主要分类创建阅读路径
            for category, papers in path_by_category.items():
                if len(papers) >= 2:  # 至少有2篇论文才构成路径
                    try:
                        category_name = parse_categories([category], 'zh-CN')[0]
                        path_name = f"{category} ({category_name})"
                    except:
                        path_name = category
                    
                    reading_path.append({
                        "path_name": path_name,
                        "papers": papers
                    })
        
        return {
            "must_read": must_read,
            "recommended": recommended,
            "interesting": interesting,
            "reading_path": reading_path
        }
    
    def generate_report(self, analysis_results, reading_recommendations=None):
        """生成分析报告"""
        print("生成分析报告...")
        
        # 1. 生成图表
        self._generate_charts(analysis_results)
        
        # 2. 生成Word文档
        self._generate_doc_report(analysis_results, reading_recommendations)
        
        # 3. 如果启用了PDF分析，则生成PDF分析报告
        if self.analyze_pdf:
            self.generate_pdf_analysis_report()
            
        # 4. 生成论文阅读推荐报告
        if reading_recommendations:
            self.generate_reading_report(reading_recommendations)
    
    def _generate_charts(self, analysis_results):
        """生成图表"""
        charts_dir = os.path.join(self.output_dir, "charts")
        os.makedirs(charts_dir, exist_ok=True)
        
        # 如果没有论文，不生成图表
        if analysis_results["total_papers"] == 0:
            return
        
        # 设置matplotlib中文字体
        self._set_matplotlib_chinese_font()
        
        # 1. 日期分布图
        plt.figure(figsize=(10, 6))
        dates = list(analysis_results["date_distribution"].keys())
        counts = list(analysis_results["date_distribution"].values())
        plt.bar(dates, counts)
        plt.title('论文日期分布')
        plt.xlabel('日期')
        plt.ylabel('论文数量')
        plt.xticks(rotation=45)
        plt.tight_layout()
        plt.savefig(os.path.join(charts_dir, 'date_distribution.png'))
        plt.close()
        
        # 2. 分类分布图（前10个分类）
        plt.figure(figsize=(12, 8))
        categories = dict(sorted(analysis_results["category_distribution"].items(), key=lambda x: x[1], reverse=True)[:10])
        cat_names = list(categories.keys())
        cat_counts = list(categories.values())
        
        # 尝试将分类代码转换为可读名称
        try:
            cat_readable = [f"{cat} ({parse_categories([cat], 'zh-CN')[0]})" for cat in cat_names]
        except:
            cat_readable = cat_names
            
        plt.barh(cat_readable, cat_counts)
        plt.title('论文分类分布（前10个分类）')
        plt.xlabel('论文数量')
        plt.ylabel('分类')
        plt.tight_layout()
        plt.savefig(os.path.join(charts_dir, 'category_distribution.png'))
        plt.close()
        
        # 3. 关键词云图
        if analysis_results["top_keywords"]:
            plt.figure(figsize=(10, 6))
            keywords = [kw for kw, _ in analysis_results["top_keywords"]]
            counts = [count for _, count in analysis_results["top_keywords"]]
            plt.barh(keywords[:15], counts[:15])
            plt.title('热门关键词（前15个）')
            plt.xlabel('出现次数')
            plt.ylabel('关键词')
            plt.tight_layout()
            plt.savefig(os.path.join(charts_dir, 'keywords.png'))
            plt.close()
        
        # 4. 趋势主题图
        if analysis_results.get("trending_topics"):
            plt.figure(figsize=(10, 6))
            topics = [topic for topic, _ in analysis_results["trending_topics"][:10]]
            topic_counts = [count for _, count in analysis_results["trending_topics"][:10]]
            
            # 使用颜色渐变增强视觉效果
            colors = plt.cm.viridis(np.linspace(0, 0.8, len(topics)))
            
            plt.barh(topics, topic_counts, color=colors)
            plt.title('研究趋势主题（前10个）')
            plt.xlabel('出现次数')
            plt.ylabel('趋势主题')
            plt.tight_layout()
            plt.savefig(os.path.join(charts_dir, 'trending_topics.png'))
            plt.close()
        
        # 5. 作者合作网络图（简化版）
        if analysis_results.get("citation_network"):
            # 找出合作最多的作者
            collaboration_counts = {}
            for author, collaborators in analysis_results["citation_network"].items():
                collaboration_counts[author] = len(collaborators)
            
            top_collaborators = sorted(collaboration_counts.items(), key=lambda x: x[1], reverse=True)[:10]
            
            if top_collaborators:
                plt.figure(figsize=(10, 6))
                authors = [author for author, _ in top_collaborators]
                collab_counts = [count for _, count in top_collaborators]
                
                # 使用颜色渐变增强视觉效果
                colors = plt.cm.cool(np.linspace(0, 0.8, len(authors)))
                
                plt.barh(authors, collab_counts, color=colors)
                plt.title('合作网络最广泛的作者（前10名）')
                plt.xlabel('合作者数量')
                plt.ylabel('作者')
                plt.tight_layout()
                plt.savefig(os.path.join(charts_dir, 'collaboration_network.png'))
                plt.close()
    
    def _set_matplotlib_chinese_font(self):
        """设置matplotlib中文字体"""
        try:
            import matplotlib.font_manager as fm
            import platform
            
            system = platform.system()
            
            # 根据操作系统设置不同的字体
            if system == 'Windows':
                # 尝试使用Windows常见中文字体
                font_paths = [
                    'C:/Windows/Fonts/simhei.ttf',  # 黑体
                    'C:/Windows/Fonts/msyh.ttc',    # 微软雅黑
                    'C:/Windows/Fonts/simsun.ttc',  # 宋体
                    'C:/Windows/Fonts/simkai.ttf'   # 楷体
                ]
                
                # 检查字体是否存在并设置
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        plt.rcParams['font.family'] = fm.FontProperties(fname=font_path).get_name()
                        print(f"已设置中文字体: {font_path}")
                        break
                
            elif system == 'Darwin':  # macOS
                # 尝试使用macOS常见中文字体
                plt.rcParams['font.family'] = ['Arial Unicode MS', 'STHeiti', 'Heiti TC', 'SimHei']
                
            elif system == 'Linux':
                # 尝试使用Linux常见中文字体
                plt.rcParams['font.family'] = ['WenQuanYi Micro Hei', 'WenQuanYi Zen Hei', 'SimHei']
            
            # 通用设置
            plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题
            
        except Exception as e:
            print(f"设置中文字体时出错: {str(e)}")
            print("将使用默认字体，中文可能无法正确显示")
    
    def _generate_doc_report(self, analysis_results, reading_recommendations=None):
        """生成Word文档报告"""
        doc = docx.Document()
        
        # 设置文档标题
        title = doc.add_heading('arXiv论文分析报告', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加日期范围
        date_range = doc.add_paragraph(f'日期范围: {self.date_from} 至 {self.date_until}')
        date_range.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加摘要
        doc.add_heading('摘要', 1)
        summary = doc.add_paragraph()
        summary.add_run(f'本报告分析了arXiv上{self.date_from}至{self.date_until}期间发布的论文。').bold = True
        summary.add_run(f'在此期间，共有{analysis_results["total_papers"]}篇符合条件的论文被收录。')
        
        if analysis_results["total_papers"] > 0:
            # 添加图表（如果有论文）
            doc.add_heading('数据可视化', 1)
            
            # 添加日期分布图
            doc.add_heading('论文日期分布', 2)
            doc.add_picture(os.path.join(self.output_dir, 'charts', 'date_distribution.png'), width=docx.shared.Inches(6))
            
            # 添加分类分布图
            doc.add_heading('论文分类分布', 2)
            doc.add_picture(os.path.join(self.output_dir, 'charts', 'category_distribution.png'), width=docx.shared.Inches(6))
            
            # 添加关键词图
            if os.path.exists(os.path.join(self.output_dir, 'charts', 'keywords.png')):
                doc.add_heading('热门关键词', 2)
                doc.add_picture(os.path.join(self.output_dir, 'charts', 'keywords.png'), width=docx.shared.Inches(6))
            
            # 添加热门作者
            doc.add_heading('热门作者', 2)
            if analysis_results["top_authors"]:
                authors_table = doc.add_table(rows=1, cols=2)
                authors_table.style = 'Table Grid'
                hdr_cells = authors_table.rows[0].cells
                hdr_cells[0].text = '作者'
                hdr_cells[1].text = '论文数量'
                
                for author, count in analysis_results["top_authors"]:
                    row_cells = authors_table.add_row().cells
                    row_cells[0].text = author
                    row_cells[1].text = str(count)
            else:
                doc.add_paragraph('没有足够的数据来分析热门作者。')
            
            # 添加趋势主题分析
            if analysis_results.get("trending_topics"):
                doc.add_heading('研究趋势分析', 2)
                trends_para = doc.add_paragraph('基于论文标题分析的热门研究趋势：')
                
                # 添加趋势主题图表
                if os.path.exists(os.path.join(self.output_dir, 'charts', 'trending_topics.png')):
                    doc.add_picture(os.path.join(self.output_dir, 'charts', 'trending_topics.png'), width=docx.shared.Inches(6))
                
                trends_table = doc.add_table(rows=1, cols=2)
                trends_table.style = 'Table Grid'
                hdr_cells = trends_table.rows[0].cells
                hdr_cells[0].text = '趋势主题'
                hdr_cells[1].text = '出现次数'
                
                for topic, count in analysis_results["trending_topics"]:
                    row_cells = trends_table.add_row().cells
                    row_cells[0].text = topic
                    row_cells[1].text = str(count)
                
                doc.add_paragraph('这些关键词反映了当前研究的热点方向和技术焦点。')
            
            # 添加作者合作网络分析
            if analysis_results.get("citation_network"):
                doc.add_heading('作者合作网络分析', 2)
                
                # 添加合作网络图表
                if os.path.exists(os.path.join(self.output_dir, 'charts', 'collaboration_network.png')):
                    doc.add_picture(os.path.join(self.output_dir, 'charts', 'collaboration_network.png'), width=docx.shared.Inches(6))
                
                # 找出合作最多的作者
                collaboration_counts = {}
                for author, collaborators in analysis_results["citation_network"].items():
                    collaboration_counts[author] = len(collaborators)
                
                top_collaborators = sorted(collaboration_counts.items(), key=lambda x: x[1], reverse=True)[:10]
                
                if top_collaborators:
                    doc.add_paragraph('以下作者在研究中展现了广泛的合作网络：')
                    
                    collab_table = doc.add_table(rows=1, cols=2)
                    collab_table.style = 'Table Grid'
                    hdr_cells = collab_table.rows[0].cells
                    hdr_cells[0].text = '作者'
                    hdr_cells[1].text = '合作者数量'
                    
                    for author, count in top_collaborators:
                        row_cells = collab_table.add_row().cells
                        row_cells[0].text = author
                        row_cells[1].text = str(count)
                    
                    # 添加一些合作者详情
                    if top_collaborators:
                        top_author = top_collaborators[0][0]
                        collaborators = analysis_results["citation_network"].get(top_author, [])
                        if collaborators:
                            doc.add_paragraph(f'作者 "{top_author}" 的主要合作者包括: {", ".join(collaborators[:5])}' + 
                                             (f' 等{len(collaborators)}人' if len(collaborators) > 5 else ''))
            
            # 添加论文阅读推荐部分
            if reading_recommendations and (reading_recommendations["must_read"] or 
                                           reading_recommendations["recommended"] or 
                                           reading_recommendations["reading_path"]):
                doc.add_heading('论文阅读推荐', 1)
                
                # 必读论文
                if reading_recommendations["must_read"]:
                    doc.add_heading('必读论文', 2)
                    for i, paper in enumerate(reading_recommendations["must_read"]):
                        p = doc.add_paragraph()
                        p.add_run(f'{i+1}. ').bold = True
                        p.add_run(f'{paper.title}').bold = True
                        p.add_run(f'\n作者: {paper.authors}')
                        p.add_run(f'\n分类: {", ".join(paper.categories)}')
                        p.add_run(f'\n摘要: {paper.abstract[:200]}...' if len(paper.abstract) > 200 else f'\n摘要: {paper.abstract}')
                        p.add_run(f'\nURL: {paper.url}')
                
                # 推荐论文
                if reading_recommendations["recommended"]:
                    doc.add_heading('推荐阅读', 2)
                    for i, paper in enumerate(reading_recommendations["recommended"][:5]):  # 只显示前5篇推荐论文
                        p = doc.add_paragraph()
                        p.add_run(f'{i+1}. ').bold = True
                        p.add_run(f'{paper.title}').bold = True
                        p.add_run(f'\n作者: {paper.authors}')
                        p.add_run(f'\n分类: {", ".join(paper.categories)}')
                        p.add_run(f'\nURL: {paper.url}')
                    
                    if len(reading_recommendations["recommended"]) > 5:
                        doc.add_paragraph(f'... 以及其他 {len(reading_recommendations["recommended"]) - 5} 篇推荐论文')
                
                # 阅读路径
                if reading_recommendations["reading_path"]:
                    doc.add_heading('学习路径推荐', 2)
                    for i, path in enumerate(reading_recommendations["reading_path"]):
                        path_para = doc.add_paragraph()
                        path_para.add_run(f'路径 {i+1}: {path["path_name"]}').bold = True
                        
                        for j, paper in enumerate(path["papers"]):
                            p = doc.add_paragraph()
                            p.style = 'List Bullet'
                            p.add_run(f'步骤 {j+1}: ').bold = True
                            p.add_run(f'{paper.title}')
                            p.add_run(f' (作者: {paper.authors})')
            
            # 添加按分类的论文列表
            doc.add_heading('按分类的论文列表', 1)
            
            for category, papers in analysis_results["papers_by_category"].items():
                try:
                    category_name = parse_categories([category], 'zh-CN')[0]
                    doc.add_heading(f'{category} ({category_name})', 2)
                except:
                    doc.add_heading(category, 2)
                
                for i, paper in enumerate(papers[:10]):  # 每个分类最多显示10篇
                    p = doc.add_paragraph()
                    p.add_run(f'{i+1}. ').bold = True
                    p.add_run(f'{paper.title}').bold = True
                    p.add_run(f'\n作者: {paper.authors}')
                    p.add_run(f'\n摘要: {paper.abstract[:200]}...' if len(paper.abstract) > 200 else f'\n摘要: {paper.abstract}')
                    p.add_run(f'\nURL: {paper.url}')
                
                if len(papers) > 10:
                    doc.add_paragraph(f'... 以及其他 {len(papers) - 10} 篇论文')
            
            # 添加按日期的论文列表
            doc.add_heading('按日期的论文列表', 1)
            
            for date_str, papers in sorted(analysis_results["papers_by_date"].items()):
                doc.add_heading(f'发布日期: {date_str}', 2)
                doc.add_paragraph(f'当日发布论文数量: {len(papers)}')
                
                # 列出该日期的前5篇论文
                for i, paper in enumerate(papers[:5]):
                    p = doc.add_paragraph()
                    p.add_run(f'{i+1}. ').bold = True
                    p.add_run(f'{paper.title}').bold = True
                    p.add_run(f'\n作者: {paper.authors}')
                    p.add_run(f'\n分类: {", ".join(paper.categories)}')
                    p.add_run(f'\nURL: {paper.url}')
                
                if len(papers) > 5:
                    doc.add_paragraph(f'... 以及其他 {len(papers) - 5} 篇论文')
        else:
            # 如果没有论文
            doc.add_paragraph('在指定的日期范围内没有找到符合条件的论文。')
        
        # 保存文档
        doc_path = os.path.join(self.output_dir, f'arxiv_report_{self.date_from}_to_{self.date_until}.docx')
        doc.save(doc_path)
        print(f"Word报告已保存到: {doc_path}")

    def call_ds_analysis(self):
        """
        调用call_ds模块中的功能来分析论文PDF并更新数据库
        """
        print("开始使用大模型分析论文PDF内容...")
        
        # 定义要分析的问题
        questions = [
            """
            论文标题和标题的中文翻译是什么？请按以下格式分别直接给出:
            英文: [英文标题]
            中文: [标题中文翻译]
            """
            ,
            """
            论文摘要和你对摘要的中文翻译是什么？请严格按以下格式分别完整地给出:
            英文: [英文摘要]
            中文: [摘要的中文翻译]
            """
            ,
            """
            文档的主要研究内容是什么？回答这个问题是可相对详细一些，并侧重关注和考虑分析是否包含有"人格、性格(personality)"的部分,前提是有对应的内容,请你按以下格式输出：
            主要内容: [文本描述]
            关于人格、性格: [文本描述]
            
            请注意，关于人格、性格的内容如果没有，则标为"无"
            """
            ,
            """
            论文使用的数据集有哪些？如果没有，则回答"无"，如果有请严格按以下格式列出：
            数据集: [数据集名称]
            数据集规模: [数据集规模]
            数据集形式: [数据集形式]
            地址: [数据集url]
            
            如果上述格式对应的项（规模、形式、地址）没有，则标为"无"，请注意数据集名称不能为"无"
            """
        ]

        col_names = [
            "question_title",
            "question_abs",
            "question_detail",
            "question_dataset"
        ]
        
        # 调用update_analyse_database函数
        try:
            # 获取日期范围内的论文
            date_from = datetime.strptime(self.date_from, "%Y-%m-%d")
            date_until = datetime.strptime(self.date_until, "%Y-%m-%d")
            
            # 首先检查数据库中是否有需要分析的论文
            conn = sqlite3.connect("papers.db")
            cursor = conn.cursor()
            
            # 检查并添加pdf列（如果不存在）
            cursor.execute("PRAGMA table_info(papers)")
            columns = [col[1] for col in cursor.fetchall()]
            
            if "pdf" not in columns:
                cursor.execute("ALTER TABLE papers ADD COLUMN pdf TEXT")
                conn.commit()
                print("已添加pdf列到数据库")
            
            # 为每个列添加列（如果不存在）
            for col in col_names:
                if col not in columns:
                    cursor.execute(f"ALTER TABLE papers ADD COLUMN {col} TEXT")
                    conn.commit()
                    print(f"已添加{col}列到数据库")
            
            # 构建日期范围条件
            date_condition = f"first_submitted_date BETWEEN '{date_from.strftime('%Y-%m-%d')}' AND '{date_until.strftime('%Y-%m-%d')}'"
            
            # 查询日期范围内有PDF但尚未分析的论文数量
            cursor.execute(f"""
                SELECT COUNT(*) 
                FROM papers 
                WHERE {date_condition}
                AND pdf IS NOT NULL
                AND (question_detail IS NULL OR question_title IS NULL)
            """)
            
            pending_count = cursor.fetchone()[0]
            
            if pending_count > 0:
                print(f"发现{pending_count}篇需要分析的论文，开始分析...")
                
                # 获取需要处理的记录
                cursor.execute(f"""
                    SELECT title, pdf, first_submitted_date
                    FROM papers 
                    WHERE {date_condition}
                    AND pdf IS NOT NULL
                    AND (question_detail IS NULL OR question_title IS NULL)
                    ORDER BY first_submitted_date DESC
                """)
                
                pending_papers = cursor.fetchall()
                conn.close()
                
                # 自定义update_analyse_database函数，使用日期范围内的论文
                from tqdm import tqdm
                
                # 创建一个新的数据库连接
                conn = sqlite3.connect("papers.db")
                cursor = conn.cursor()
                
                # 添加进度条包装迭代对象
                for title, pdf, first_submitted_date in tqdm(pending_papers, desc="分析论文中", unit="篇"):
                    try:
                        tqdm.write(f"正在处理：{pdf}")  # 使用 tqdm.write 避免与进度条冲突
                        pdf_path = pdf
                        full_text = load_pdf_content(pdf_path)
                        if full_text:
                            results = chat_v3(full_text, questions)
                        else:
                            tqdm.write(f"无法加载PDF内容: {pdf}")
                            continue
                    except Exception as e:
                        tqdm.write(f"处理PDF时出错 [{pdf}]: {str(e)}")
                        continue

                    try:
                        for i in range(len(results)):
                            # 更新数据库记录
                            cursor.execute(
                                f"UPDATE papers SET {col_names[i]} = ? WHERE title = ?",
                                (f"{results[i]}", title)
                            )
                            conn.commit()
                        tqdm.write(f"成功更新：{pdf_path}")

                    except Exception as e:
                        tqdm.write(f"更新数据库失败 [{pdf}]: {str(e)}")
                        conn.rollback()
                        continue
                
                conn.close()
                print("论文分析完成！")
            else:
                print("没有找到需要分析的论文PDF，跳过分析步骤")
                conn.close()
                
        except Exception as e:
            print(f"分析论文时发生错误: {str(e)}")
            import traceback
            traceback.print_exc()

    def download_pdfs(self, papers):
        """
        下载论文的PDF文件
        
        Args:
            papers: 论文列表
        """
        print("开始下载论文PDF...")
        
        # 创建PDF保存目录
        pdf_dir = os.path.join(self.output_dir, "pdfs")
        os.makedirs(pdf_dir, exist_ok=True)
        
        # 连接到数据库
        conn = sqlite3.connect("papers.db")
        cursor = conn.cursor()
        
        # 检查pdf列是否存在
        cursor.execute("PRAGMA table_info(papers)")
        columns = [col[1] for col in cursor.fetchall()]
        
        if "pdf" not in columns:
            cursor.execute("ALTER TABLE papers ADD COLUMN pdf TEXT")
            conn.commit()
            print("已添加pdf列到数据库")
        
        # 获取已经下载的PDF列表
        cursor.execute("SELECT title, pdf FROM papers WHERE pdf IS NOT NULL")
        existing_pdfs = {title: pdf_path for title, pdf_path in cursor.fetchall()}
        
        # 下载PDF
        download_count = 0
        skip_count = 0
        error_count = 0
        
        for paper in tqdm(papers, desc="下载PDF", unit="篇"):
            # 如果已经下载过，则跳过
            if paper.title in existing_pdfs and os.path.exists(existing_pdfs[paper.title]):
                skip_count += 1
                continue
            
            # 构建PDF URL
            if paper.url:
                # 将arxiv.org/abs/XXXX.XXXXX 转换为 arxiv.org/pdf/XXXX.XXXXX.pdf
                parsed_url = urlparse(paper.url)
                if parsed_url.netloc == 'arxiv.org' and parsed_url.path.startswith('/abs/'):
                    paper_id = parsed_url.path.replace('/abs/', '')
                    pdf_url = f"https://arxiv.org/pdf/{paper_id}.pdf"
                    
                    try:
                        # 下载PDF
                        response = requests.get(pdf_url, stream=True)
                        if response.status_code == 200:
                            # 构建保存路径
                            safe_title = "".join([c if c.isalnum() or c in [' ', '.', '-', '_'] else '_' for c in paper.title])
                            safe_title = safe_title[:100]  # 限制文件名长度
                            pdf_filename = f"{paper_id}_{safe_title}.pdf"
                            pdf_path = os.path.join(pdf_dir, pdf_filename)
                            
                            # 保存PDF
                            with open(pdf_path, 'wb') as f:
                                for chunk in response.iter_content(chunk_size=8192):
                                    if chunk:
                                        f.write(chunk)
                            
                            # 更新数据库
                            cursor.execute(
                                "UPDATE papers SET pdf = ? WHERE title = ?",
                                (pdf_path, paper.title)
                            )
                            conn.commit()
                            download_count += 1
                        else:
                            tqdm.write(f"下载失败 [{paper.title}]: HTTP状态码 {response.status_code}")
                            error_count += 1
                    except Exception as e:
                        tqdm.write(f"下载出错 [{paper.title}]: {str(e)}")
                        error_count += 1
                else:
                    tqdm.write(f"不支持的URL格式 [{paper.url}]")
                    error_count += 1
            else:
                tqdm.write(f"没有URL [{paper.title}]")
                error_count += 1
        
        conn.close()
        print(f"PDF下载完成！成功: {download_count}, 跳过: {skip_count}, 失败: {error_count}")

    def generate_pdf_analysis_report(self):
        """生成PDF分析报告"""
        print("生成PDF分析报告...")
        
        # 连接到数据库
        conn = sqlite3.connect("papers.db")
        cursor = conn.cursor()
        
        # 获取日期范围
        date_from = datetime.strptime(self.date_from, "%Y-%m-%d")
        date_until = datetime.strptime(self.date_until, "%Y-%m-%d")
        date_condition = f"first_submitted_date BETWEEN '{date_from.strftime('%Y-%m-%d')}' AND '{date_until.strftime('%Y-%m-%d')}'"
        
        # 查询已分析的论文
        cursor.execute(f"""
            SELECT title, question_title, question_abs, question_detail, question_dataset
            FROM papers 
            WHERE {date_condition}
            AND question_detail IS NOT NULL
            ORDER BY first_submitted_date DESC
        """)
        
        analyzed_papers = cursor.fetchall()
        conn.close()
        
        if not analyzed_papers:
            print("没有找到已分析的论文，跳过生成PDF分析报告")
            return
        
        # 创建Word文档
        doc = docx.Document()
        
        # 设置文档标题
        title = doc.add_heading('arXiv论文PDF分析报告', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加日期范围
        date_range = doc.add_paragraph(f'日期范围: {self.date_from} 至 {self.date_until}')
        date_range.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加摘要
        doc.add_heading('摘要', 1)
        summary = doc.add_paragraph()
        summary.add_run(f'本报告分析了arXiv上{self.date_from}至{self.date_until}期间发布的论文PDF内容。').bold = True
        summary.add_run(f'在此期间，共有{len(analyzed_papers)}篇论文被深入分析。')
        
        # 添加论文分析结果
        doc.add_heading('论文分析结果', 1)
        
        for i, paper in enumerate(analyzed_papers):
            title, question_title, question_abs, question_detail, question_dataset = paper
            
            # 添加论文标题
            doc.add_heading(f'{i+1}. {title}', 2)
            
            # 添加标题翻译
            if question_title:
                doc.add_heading('标题翻译', 3)
                doc.add_paragraph(question_title)
            
            # 添加摘要翻译
            if question_abs:
                doc.add_heading('摘要翻译', 3)
                doc.add_paragraph(question_abs)
            
            # 添加研究内容
            if question_detail:
                doc.add_heading('研究内容', 3)
                doc.add_paragraph(question_detail)
            
            # 添加数据集信息
            if question_dataset and question_dataset != "无":
                doc.add_heading('数据集信息', 3)
                doc.add_paragraph(question_dataset)
            
            # 添加分隔线
            if i < len(analyzed_papers) - 1:
                doc.add_paragraph('_' * 50)
        
        # 保存文档
        doc_path = os.path.join(self.output_dir, f'arxiv_pdf_analysis_{self.date_from}_to_{self.date_until}.docx')
        doc.save(doc_path)
        print(f"PDF分析报告已保存到: {doc_path}")

    def generate_reading_report(self, reading_recommendations):
        """生成论文阅读推荐报告"""
        print("生成论文阅读推荐报告...")
        
        if not (reading_recommendations["must_read"] or 
                reading_recommendations["recommended"] or 
                reading_recommendations["interesting"]):
            print("没有阅读推荐，跳过生成阅读推荐报告")
            return
        
        # 创建Word文档
        doc = docx.Document()
        
        # 设置文档标题
        title = doc.add_heading('arXiv论文阅读推荐', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加日期范围
        date_range = doc.add_paragraph(f'日期范围: {self.date_from} 至 {self.date_until}')
        date_range.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加摘要
        doc.add_heading('摘要', 1)
        summary = doc.add_paragraph()
        summary.add_run(f'本报告为arXiv上{self.date_from}至{self.date_until}期间发布的论文提供阅读推荐。').bold = True
        total_papers = len(reading_recommendations["must_read"]) + len(reading_recommendations["recommended"]) + len(reading_recommendations["interesting"])
        summary.add_run(f'共有{total_papers}篇论文被推荐阅读，其中必读论文{len(reading_recommendations["must_read"])}篇，推荐阅读{len(reading_recommendations["recommended"])}篇，有趣论文{len(reading_recommendations["interesting"])}篇。')
        
        # 必读论文
        if reading_recommendations["must_read"]:
            doc.add_heading('必读论文', 1)
            doc.add_paragraph('以下论文是本期最重要的研究成果，强烈建议阅读：')
            
            for i, paper in enumerate(reading_recommendations["must_read"]):
                doc.add_heading(f'{i+1}. {paper.title}', 2)
                
                # 添加作者
                authors_para = doc.add_paragraph()
                authors_para.add_run('作者: ').bold = True
                authors_para.add_run(paper.authors)
                
                # 添加分类
                categories_para = doc.add_paragraph()
                categories_para.add_run('分类: ').bold = True
                categories_para.add_run(', '.join(paper.categories))
                
                # 添加摘要
                abstract_para = doc.add_paragraph()
                abstract_para.add_run('摘要: ').bold = True
                abstract_para.add_run(paper.abstract)
                
                # 添加URL
                url_para = doc.add_paragraph()
                url_para.add_run('URL: ').bold = True
                url_para.add_run(paper.url)
                
                # 添加分隔线
                if i < len(reading_recommendations["must_read"]) - 1:
                    doc.add_paragraph('_' * 50)
        
        # 推荐阅读
        if reading_recommendations["recommended"]:
            doc.add_heading('推荐阅读', 1)
            doc.add_paragraph('以下论文值得关注，建议在有时间时阅读：')
            
            for i, paper in enumerate(reading_recommendations["recommended"]):
                doc.add_heading(f'{i+1}. {paper.title}', 2)
                
                # 添加作者
                authors_para = doc.add_paragraph()
                authors_para.add_run('作者: ').bold = True
                authors_para.add_run(paper.authors)
                
                # 添加分类
                categories_para = doc.add_paragraph()
                categories_para.add_run('分类: ').bold = True
                categories_para.add_run(', '.join(paper.categories))
                
                # 添加摘要（简短版本）
                abstract_para = doc.add_paragraph()
                abstract_para.add_run('摘要: ').bold = True
                abstract_para.add_run(paper.abstract[:200] + ('...' if len(paper.abstract) > 200 else ''))
                
                # 添加URL
                url_para = doc.add_paragraph()
                url_para.add_run('URL: ').bold = True
                url_para.add_run(paper.url)
                
                # 添加分隔线
                if i < len(reading_recommendations["recommended"]) - 1:
                    doc.add_paragraph('_' * 50)
        
        # 有趣论文
        if reading_recommendations["interesting"]:
            doc.add_heading('有趣论文', 1)
            doc.add_paragraph('以下论文可能包含有趣的研究方向或结果：')
            
            # 创建表格
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # 添加表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '标题'
            hdr_cells[1].text = '作者'
            hdr_cells[2].text = 'URL'
            
            # 添加论文
            for paper in reading_recommendations["interesting"]:
                row_cells = table.add_row().cells
                row_cells[0].text = paper.title
                row_cells[1].text = paper.authors
                row_cells[2].text = paper.url
        
        # 学习路径
        if reading_recommendations["reading_path"]:
            doc.add_heading('学习路径推荐', 1)
            doc.add_paragraph('以下是基于论文主题和难度的学习路径推荐，按照路径顺序阅读可以更好地理解相关研究领域：')
            
            for i, path in enumerate(reading_recommendations["reading_path"]):
                doc.add_heading(f'路径 {i+1}: {path["path_name"]}', 2)
                
                for j, paper in enumerate(path["papers"]):
                    step_para = doc.add_paragraph(style='List Bullet')
                    step_para.add_run(f'步骤 {j+1}: ').bold = True
                    step_para.add_run(f'{paper.title}')
                    step_para.add_run(f' (作者: {paper.authors})')
                    step_para.add_run(f'\nURL: {paper.url}')
        
        # 保存文档
        doc_path = os.path.join(self.output_dir, f'arxiv_reading_recommendations_{self.date_from}_to_{self.date_until}.docx')
        doc.save(doc_path)
        print(f"阅读推荐报告已保存到: {doc_path}")

async def main():
    try:
        # 添加命令行参数解析
        parser = argparse.ArgumentParser(description='arXiv论文爬取与分析工作流')
        parser.add_argument('--date-from', type=str, help='开始日期 (YYYY-MM格式)')
        parser.add_argument('--date-until', type=str, help='结束日期 (YYYY-MM格式)')
        parser.add_argument('--analyze-pdf', action='store_true', help='是否使用大模型分析PDF内容')
        args = parser.parse_args()
        
        # 注意：arXiv API只接受年月格式的日期，不接受具体到日的日期
        # 使用命令行参数或默认值
        date_from = args.date_from if args.date_from else "2025-02-01"
        date_until = args.date_until if args.date_until else "2025-02-28"
        
        # 确保日期格式正确（只保留年月部分）
        if len(date_from) > 7:  # 如果格式是YYYY-MM-DD
            date_from = date_from[:7] + "-01"  # 转换为YYYY-MM-01
        if len(date_until) > 7:  # 如果格式是YYYY-MM-DD
            # 获取月末日期
            year, month = map(int, date_until[:7].split('-'))
            if month == 12:
                next_month = datetime(year + 1, 1, 1) - timedelta(days=1)
            else:
                next_month = datetime(year, month + 1, 1) - timedelta(days=1)
            date_until = next_month.strftime("%Y-%m-%d")
        
        print(f"注意：arXiv API只接受年月格式的日期，实际搜索范围将是从{date_from[:7]}月到{date_until[:7]}月")
        
        # 创建工作流实例
        workflow = ArxivWorkflow(
            category_whitelist=["cs.AI", "cs.CL", "cs.LG", "cs.CV"],
            optional_keywords=["LLM", "GPT", "AI", "language model", "deep learning", 
                              "transformer", "neural network", "machine learning"],
            trans_to=None,  # 暂时不启用翻译
            date_from=date_from,
            date_until=date_until,
            analyze_pdf=args.analyze_pdf  # 传递命令行参数
        )
        
        # 打印实际使用的日期范围，用于调试
        print(f"实际爬取日期范围: {workflow.date_from} 到 {workflow.date_until}")
        
        # 运行工作流
        await workflow.run()
        
    except Exception as e:
        print(f"工作流执行过程中发生错误: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    asyncio.run(main()) 